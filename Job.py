#!/usr/bin/env python
# coding: utf-8

# In[ ]:


get_ipython().system('pip install -qU "langchain[groq]"')
get_ipython().system('pip install -qU langchain-text-splitters')
get_ipython().system('pip install -qU langchain-chroma')
get_ipython().system('pip install -U langchain-community')
get_ipython().system('pip install assemblyai')
get_ipython().system('pip install python-docx')
get_ipython().system('pip install streamlit-audiorecorder')
get_ipython().system('pip install streamlit')


# In[ ]:


#key-grok =  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
#key-huggingface = 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
#assemblyai = cce30862c26442f09e577fbe639108ed


# In[ ]:


import getpass
import os
from groq import Groq
from langchain.embeddings import SentenceTransformerEmbeddings
from sentence_transformers import SentenceTransformer
from langchain.chat_models import init_chat_model
from langchain.embeddings import SentenceTransformerEmbeddings
from sentence_transformers import SentenceTransformer
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter


# Initialize the model
model_name = 'all-MiniLM-L6-v2' # assign the model name to a variable

# Create an embedding instance
embedding = SentenceTransformerEmbeddings(model_name=model_name) # use model_name keyword argument

vector_store = Chroma(
    collection_name="example_collection",
    embedding_function=embedding,
    persist_directory="./chroma_langchain_db",  # Where to save data locally, remove if not necessary
)

os.environ["GROQ_API_KEY"] = "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"

if not os.environ.get("GROQ_API_KEY"):
  os.environ["GROQ_API_KEY"] = getpass.getpass("Enter API key for Groq: ")


model = init_chat_model("llama3-8b-8192", model_provider="groq")

client = Groq()


text_splitter = RecursiveCharacterTextSplitter(
    # Set a really small chunk size, just to show.
    chunk_size=1000,  # Increased chunk size
    chunk_overlap=50,  # Reduced chunk overlap
    length_function=len,
    is_separator_regex=False,

)




# Create an embedding instance



# In[ ]:


import assemblyai as aai
from langchain_core.documents import Document

# Set AssemblyAI API key (get from assembly.ai after signup)
aai.settings.api_key = "cce30862c26442f09e577fbe639108ed"

# Configure transcription with diarization
config = aai.TranscriptionConfig(speaker_labels=True)
transcriber = aai.Transcriber()

# Transcribe audio file (local or URL)
audio_file = "audio.mp3"  # Replace with your audio file
transcript = transcriber.transcribe(audio_file, config=config)

# Create LangChain Document with speaker-tagged text
text_with_speakers = ""
for utterance in transcript.utterances:
    text_with_speakers += f"Speaker {utterance.speaker} [{utterance.start/1000:.1f}s - {utterance.end/1000:.1f}s]: {utterance.text}\n"

doc = Document(page_content=text_with_speakers)

# Print the result
print(doc.page_content)


# In[ ]:


print(transcript.utterances[40].text)
temp=model.invoke("""Perform sentiment ananlysis and
Classify the following text on a scale of 1 to 10. 1 being bad and
10 being the best. Give a 1 word answer: """+transcript.utterances[29].text)
print(temp.content)


# In[ ]:


import os
from docx import Document

def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def read_all_docx_in_folder(folder_path):
    all_texts = {}
    for filename in os.listdir(folder_path):
        print(filename)
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)
            try:
                text = read_docx(file_path)
                all_texts[filename] = text
            except Exception as e:
                print(f"Failed to read {filename}: {e}")
    return all_texts

folder_path = "RagDocs"
documents = read_all_docx_in_folder(folder_path)

for filename, content in documents.items():
#    print(content)
    state_of_the_union =content
    texts = text_splitter.create_documents([state_of_the_union])
    texts_content = [t.page_content for t in texts]
    vector_store.add_texts(texts=texts_content)


# In[ ]:


query = "I have a Stb issue how do i solve it"
results = vector_store.similarity_search(query, k=20)
context=" "
for r in results:
    context=context+r.page_content
#    print(f"Match: {r.page_content}")
ans=model.invoke(" Answer the question based on the following context:"+ context+"Answer the question based on the above context:"+query)
print(ans.content)
#print(context)


# In[ ]:


from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
# Instead of HuggingFacePipeline, call the pipeline directly
from transformers import pipeline

# Initialize the sentiment analysis pipeline
sentiment_pipeline = pipeline("text-classification", model="tabularisai/multilingual-sentiment-analysis")

# Create a LangChain prompt template (optional, you can skip this and call pipeline directly)
prompt_template = PromptTemplate(
    input_variables=["text"],
    template="Classify the sentiment of the following text: {text}"
)

# Create a function to use the pipeline for prediction
def predict_sentiment(text):
    result = sentiment_pipeline(text)[0]  # Get the prediction from the pipeline
    return result['label']  # Extract the sentiment label

# Example usage
text = transcript.utterances[25].text
# Instead of llm_chain.run, call predict_sentiment
result = predict_sentiment(text)
print(result)
print(text)


# In[ ]:


get_ipython().system('streamlit run app.py & npx localtunnel --port 8501')


# In[ ]:


temp=model.invoke("""Perform sentiment ananlysis and
Classify the following text on a scale of 1 to 10. 1 being bad and
10 being the best. Give a 1 word answer: Pakistan is best""")
print(temp.content)

